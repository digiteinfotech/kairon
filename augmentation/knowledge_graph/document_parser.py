from operator import itemgetter
from collections import deque, Counter
import re
import pymupdf
from docx import Document


class DocumentParser:

    @staticmethod
    def fonts(doc, granularity=False):
        """
        Extracts fonts and their usage in PDF documents

        :param doc: PDF document to iterate through
        :param granularity: also use 'font', 'flags' and 'color' to discriminate text
        :return: list of most used fonts sorted by count, font style information
        """
        styles = {}
        font_counts = {}

        for page in doc:
            blocks = page.get_text("dict")["blocks"]

            for b in blocks:  # iterate through the text blocks
                if b['type'] == 0:  # block contains text
                    for l in b["lines"]:  # iterate through the text lines
                        for s in l["spans"]:  # iterate through the text spans
                            if granularity:
                                identifier = "{0}_{1}_{2}_{3}".format(s['size'], s['flags'], s['font'], s['color'])
                                styles[identifier] = {'size': s['size'], 'flags': s['flags'], 'font': s['font'],
                                                      'color': s['color']}
                            else:
                                identifier = "{0}".format(s['size'])
                                styles[identifier] = {'size': s['size'], 'font': s['font']}

                            font_counts[identifier] = font_counts.get(identifier, 0) + 1  # count the fonts usage

        font_counts = sorted(font_counts.items(), key=itemgetter(1), reverse=True)

        if not font_counts:
            raise ValueError("Zero discriminating fonts found!")

        return font_counts, styles

    @staticmethod
    def font_tags(font_counts, styles):
        """
        Returns dictionary with font sizes as keys and tags as value

        :param font_counts: (font_size, count) for all fonts occuring in document
        :param styles: all styles found in the document
        :return: all element tags based on font-sizes
        """
        p_style = styles[font_counts[0][0]]  # get style for most used font by count (paragraph)
        p_size = p_style['size']  # get the paragraph's size

        # sorting the font sizes high to low, so that we can append the right integer to each tag
        font_sizes = []
        for (font_size, count) in font_counts:
            font_sizes.append(float(font_size))
        font_sizes.sort(reverse=True)

        # aggregating the tags for each font size
        idx = 0
        size_tag = {}
        for size in font_sizes:
            idx += 1
            if size == p_size:
                idx = 0
                size_tag[size] = '<p>'
            if size > p_size:
                size_tag[size] = '<h{0}>'.format(idx)
            elif size < p_size:
                size_tag[size] = '<s{0}>'.format(idx)

        return size_tag

    @staticmethod
    def headers_paragraphs(doc, size_tag):
        """
        Scrapes headers & paragraphs from PDF and return texts with element tags

        :param doc: PDF document to iterate through
        :param size_tag: textual element tags for each size
        :return: list of texts with pre-prended element tags
        """
        header_para = []  # list with headers and paragraphs
        first = True  # boolean operator for first header
        previous_s = {}  # previous span
        for page in doc:
            blocks = page.get_text("dict")["blocks"]
            for b in blocks:  # iterate through the text blocks
                if b['type'] == 0:  # this block contains text

                    # REMEMBER: multiple fonts and sizes are possible IN one block

                    block_string = ""  # text found in block
                    for l in b["lines"]:  # iterate through the text lines

                        for s in l["spans"]:  # iterate through the text spans

                            if s['text'].strip():  # removing whitespaces:
                                if first:
                                    previous_s = s
                                    first = False
                                    block_string = size_tag[s['size']] + " " + s['text']
                                else:
                                    if s['size'] == previous_s['size']:

                                        if block_string and all((c == "|") for c in block_string):
                                            # block_string only contains pipes
                                            block_string = size_tag[s['size']] + " " + s['text']
                                        if block_string == "":
                                            # new block has started, so append size tag
                                            block_string = size_tag[s['size']] + " " + s['text']
                                        else:  # in the same block, so concatenate strings
                                            block_string += " " + s['text']

                                    else:
                                        header_para.append(block_string)
                                        block_string = size_tag[s['size']] + " " + s['text']

                                    previous_s = s

                        # new block started, indicating with a pipe
                        block_string += "|"

                    header_para.append(block_string)

        return header_para

    @staticmethod
    def pdf(path):
        """
        Accepts the path of the pdf file and processes it

        :param path: Path of the pdf file
        :return: list of sentences and dictionary structure of the document
        """
        document = path
        doc = pymupdf.open(document)

        # get the allowed font sizes
        font_counts, styles = DocumentParser.fonts(doc, granularity=False)
        allowed_sizes = []
        para = float(font_counts[0][0])
        for element in font_counts:
            if float(element[0]) >= para:
                allowed_sizes.append(float(element[0]))
        allowed_sizes.sort(reverse=True)

        # get tag to size dictionary
        size_dict = {}
        size_dict[allowed_sizes[-1]] = "<p>"
        for i in range(len(allowed_sizes) - 1):
            size_dict[allowed_sizes[i]] = "<h" + str(i + 1) + ">"
        no_diff_fonts = len(allowed_sizes)
        highestSize = no_diff_fonts
        tagtosize = {}
        for i in range(no_diff_fonts):
            tagtosize[size_dict[allowed_sizes[i]][1:-1]] = highestSize
            highestSize -= 1

        # get list of strings with tags and list of priority by number
        size_tag = DocumentParser.font_tags(font_counts, styles)
        elements = DocumentParser.headers_paragraphs(doc, size_tag)
        elements = [i.replace('|', '') for i in elements]
        elements = [i for i in elements if len(i.strip()) > 0]
        elements2 = [i for i in elements if not i.replace(i[i.find("<"):i.find(">") + 1], '').strip().isdigit()]
        qw = [item for item, count in Counter(elements2).items() if count > 5 and '<h' not in item]
        final_list = [item for item in elements2 if item not in qw]
        final_list = [item for item in final_list if not '<s' in item]
        doc_list = final_list
        docsize_list = []
        for string in doc_list:
            tag = string[string.find("<") + 1:string.find(">")]
            docsize_list.append(tagtosize[tag])

        # remove consecutive duplicates
        reducedoclist = []
        preValue = -1
        for element_n in docsize_list:
            value = element_n
            if value != 1:
                reducedoclist.append(value)
            else:
                if preValue != value:
                    reducedoclist.append(value)
            preValue = value

        # merge continuous tags
        newlist = []
        sizelistReduced = []
        string = doc_list[0]
        tag = string[string.find("<") + 1:string.find(">")]
        for element in doc_list[1:]:
            tag1 = element[element.find("<") + 1:element.find(">")]
            if tag1 != tag:
                newlist.append(string)
                sizelistReduced.append(tagtosize[tag])
                string = element
                tag = tag1
            else:
                element = re.sub('<' + tag1 + '>', '', element)
                if string.strip()[-1].isalpha() or string.strip()[-1].isdigit():
                    string = string + ', ' + element
                else:
                    string = string + " " + element
        newlist.append(string)
        sizelistReduced.append(tagtosize[tag])

        # order the strings based on tags
        thislist = sizelistReduced
        helperstack = deque()
        helperid = deque()
        index = 1
        arrSize = len(thislist)
        treedict = {}
        helperstack.append(thislist[arrSize - 1])
        helperid.append(arrSize - 1)
        while helperstack:
            value = thislist[arrSize - index - 1]
            if value > helperstack[-1]:
                treedict[arrSize - index - 1] = []
                while helperstack and (value > helperstack[-1]):
                    helperstack.pop()
                    treedict[arrSize - index - 1].append(helperid.pop())
            helperstack.append(value)
            helperid.append(arrSize - index - 1)
            index += 1
            if index >= arrSize:
                break

        return treedict, newlist

    @staticmethod
    def docx(path):
        """
        Accepts the path of the docx file and processes it

        :param path: Path of the docx file
        :return: list of sentences and dictionary structure of the document
        """
        doc = Document(path)
        size_list = [p.style.font.size for p in doc.paragraphs]

        # get the allowed font sizes
        A = Counter(size_list).most_common()
        para = A[0][0]
        allowed_sizes = []
        for s in A:
            if str(s[0]).isdigit() and s[0] >= para:
                allowed_sizes.append(s[0])
        allowed_sizes.sort(reverse=True)

        # get tag to size dictionary
        size_dict = {}
        sizeorder_dict = {}
        size_dict[allowed_sizes[-1]] = "<p>"
        for i in range(len(allowed_sizes) - 1):
            size_dict[allowed_sizes[i]] = "<h" + str(i) + ">"
        no_diff_fonts = len(allowed_sizes)
        highestSize = no_diff_fonts
        tagtosize = {}
        for i in range(no_diff_fonts):
            sizeorder_dict[allowed_sizes[i]] = highestSize
            tagtosize[size_dict[allowed_sizes[i]][1:-1]] = highestSize
            highestSize -= 1

        # get list of strings with tags and list of priority by number
        doc_list = []
        docsize_list = []
        for p in doc.paragraphs:
            size = p.style.font.size
            if size in size_dict:
                text = p.text.strip()
                if text != '':
                    tag = size_dict[size]
                    doc_list.append(tag + " " + text)
                    docsize_list.append(sizeorder_dict[size])

        # remove consecutive duplicates
        reducedoclist = []
        preValue = -1
        for element_n in docsize_list:
            value = element_n
            if value != 1:
                reducedoclist.append(value)
            else:
                if preValue != value:
                    reducedoclist.append(value)
            preValue = value

        # merge continuous tags
        newlist = []
        sizelistReduced = []
        string = doc_list[0]
        tag = string[string.find("<") + 1:string.find(">")]
        for element in doc_list[1:]:
            tag1 = element[element.find("<") + 1:element.find(">")]
            if tag1 != tag:
                newlist.append(string)
                sizelistReduced.append(tagtosize[tag])
                string = element
                tag = tag1
            else:
                element = re.sub('<' + tag1 + '>', '', element)
                if string.strip()[-1].isalpha() or string.strip()[-1].isdigit():
                    string = string + ', ' + element
                else:
                    string = string + " " + element
        newlist.append(string)
        sizelistReduced.append(tagtosize[tag])

        # order the strings based on tags
        thislist = sizelistReduced
        helperstack = deque()
        helperid = deque()
        index = 1
        arrSize = len(thislist)
        tree_struct = {}
        helperstack.append(thislist[arrSize - 1])
        helperid.append(arrSize - 1)
        while helperstack:
            value = thislist[arrSize - index - 1]
            if value > helperstack[-1]:
                tree_struct[arrSize - index - 1] = []
                while helperstack and (value > helperstack[-1]):
                    helperstack.pop()
                    tree_struct[arrSize - index - 1].append(helperid.pop())
            helperstack.append(value)
            helperid.append(arrSize - index - 1)
            index += 1
            if index >= arrSize:
                break

        return tree_struct, newlist

    @staticmethod
    def parse(path):
        if path.lower().endswith('.pdf'):
            tree_struct, new_list = DocumentParser.pdf(path)
        else:
            tree_struct, new_list = DocumentParser.docx(path)

        return tree_struct, new_list